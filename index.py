import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import re
from openpyxl import Workbook
import pymorphy2
import shutil

class GUITestApp:
    def __init__(self, root):
        self.root = root
        self.root.title("GUI Test")
        self.root.geometry("1200x600")
        self.instruction = """Паспорт - Passport  
Кем признан - Called  
Должность - Lvlrank  
Дать поступление на военную службу - Receiptdate  
Дата получения травмы -  Injury  
Дата предоставление документа - Datec  
Фио - Fio  
Воинская часть - Part  
Звание - Rank  
Дата рождение - Agedate  
Дата поступление - Receiveddate  
Жалобы - Pity  
Анамнез - Information  
Объективный статус - Condition  
Сердечно сосудистая система - Vessels  
Дыхательная система - Breath  
Система органов пищеварения - Digestion  
Мочеполовая система - Urinary  
Status Localis - Common  
Данные лабораторных методов исследования - Laboratory  
Данные инструментальных методов исследования - Tools  
Получал терапию - Therapy  
(Прикрутить склонение) Диагноз - Diagnosismakeup  
И диагноз без склонения - Diagnosis 
Номер № Numb """
        self.template_folder = None
        self.check_all_var = tk.BooleanVar(value=False)
        self.data = []
        self.setup_ui()

    def setup_ui(self):
        self.create_menu_frame()
        self.create_checkboxes_frame()
        self.create_table_frame()
        self.create_start_button()

    def create_menu_frame(self):
        self.menu_frame = tk.Frame(self.root)
        self.menu_frame.grid(row=0, column=0, columnspan=3, sticky='ew', padx=10, pady=10)

        self.file_name_label = tk.Label(self.menu_frame, text="No file selected")
        self.file_name_label.pack(side=tk.LEFT, padx=5)

        self.open_file_button = tk.Button(self.menu_frame, text="Open File", command=self.open_file)
        self.open_file_button.pack(side=tk.LEFT, padx=5)

        self.open_folder_button = tk.Button(self.menu_frame, text="Open Template Folder",
                                            command=self.open_template_folder)
        self.open_folder_button.pack(side=tk.LEFT, padx=5)

    def create_checkboxes_frame(self):
        self.checkboxes_frame = tk.Frame(self.root)
        self.checkboxes_frame.grid(row=1, column=0, sticky='ns', padx=10, pady=10)

        self.button_check_all = tk.Button(self.root, text="Выбрать всё/Отменить всё", command=self.checkall)
        self.button_check_all.grid(row=2, column=0, sticky='ns', padx=10, pady=10)

        self.checkboxes = []

    def create_table_frame(self):
        self.table_frame = tk.Frame(self.root)
        self.table_frame.grid(row=1, column=2, sticky='nsew', padx=10, pady=10)

        self.table = ttk.Treeview(self.table_frame, columns=('col1', 'col2'), show='headings')
        self.table.heading('col1', text='Column 1')
        self.table.heading('col2', text='Column 2')
        self.table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.scrollbar = ttk.Scrollbar(self.table_frame, orient=tk.VERTICAL, command=self.table.yview)
        self.table.configure(yscroll=self.scrollbar.set)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.table.bind('<Double-1>', self.edit_cell)
        self.table.bind('<Return>', self.save_edited_cell)

        self.button_frame = tk.Frame(self.root)
        self.button_frame.grid(row=1, column=3, sticky='ns', padx=10, pady=10)

        # Guide label with self.instruction text
        self.guide = tk.Label(self.root, text=self.instruction)
        self.guide.grid(row=1, column=1, sticky='ns', padx=10, pady=10)

        self.add_row_button = tk.Button(self.button_frame, text="+", command=self.add_row)
        self.add_row_button.pack(pady=5)

        self.remove_row_button = tk.Button(self.button_frame, text="-", command=self.remove_row)
        self.remove_row_button.pack(pady=5)

        self.save_table_button = tk.Button(self.button_frame, text="Save Table", command=self.save_table)
        self.save_table_button.pack(pady=5)

        edit_row = ['Passport', 'Called', 'Lvlrank', 'Receiptdate', 'Injury', 'Datec']
        for value in edit_row:
            self.table.insert('', 'end', values=(value, ""))

    def add_row(self):
        self.data.append(["", ""])
        self.table.insert('', 'end', values=("", ""))

    def remove_row(self):
        selected_item = self.table.selection()
        if selected_item:
            self.table.delete(selected_item)
            for i in selected_item:
                self.data.pop(self.table.index(i))

    def edit_cell(self, event):
        row_id = self.table.identify_row(event.y)
        column_id = self.table.identify_column(event.x)
        if row_id and column_id:
            self.table.focus(row_id)
            self.table.selection_set(row_id)

            x, y, width, height = self.table.bbox(row_id, column_id)
            entry = tk.Entry(self.table)
            entry.place(x=x, y=y, width=width, height=height, anchor='nw')
            entry.insert(0, self.table.item(row_id, 'values')[int(column_id[1]) - 1])

            entry.bind('<Return>', lambda e: self.save_edited_cell(e, entry, row_id, column_id))
            entry.bind('<FocusOut>', lambda e: entry.destroy())

            entry.focus_set()

    def save_edited_cell(self, event, entry=None, row_id=None, column_id=None):
        if entry:
            new_value = entry.get()
            values = list(self.table.item(row_id, 'values'))
            values[int(column_id[1]) - 1] = new_value
            self.table.item(row_id, values=values)
            entry.destroy()

            # Update the data array with the new value
            updated = False
            for i, item in enumerate(self.data):
                if item[0] == row_id:  # Assuming row_id is the first element in each sublist
                    self.data[i][int(column_id[1]) - 1] = new_value
                    updated = True
                    break



            if not updated:
                # If row_id doesn't exist in self.data, append a new row
                self.data.append([row_id] + [""] * (int(column_id[1]) - 1) + [new_value])
                print("not update")

    def save_table(self):
        # Implement saving the self.data array to a file or database as needed
        print("Data saved:", self.data)

    def create_start_button(self):
        # Создаем первую кнопку
        self.start_button = tk.Button(self.root, text="Отрисовать данные", command=self.edit_process)
        self.start_button.grid(row=2, column=1, columnspan=2, pady=(0, 5))  # pady=(0, 5) добавляет отступ снизу

        # Создаем вторую кнопку
        self.process_button = tk.Button(self.root, text="Начать процесс", command=self.start_process)
        self.process_button.grid(row=3, column=1, columnspan=2,
                                 pady=(5, 10))  # pady=(5, 10) добавляет отступ сверху и снизу

    def open_file(self):
        self.file_path = filedialog.askopenfilename()
        if self.file_path:
            self.file_name_label.config(text=os.path.basename(self.file_path))

    def open_template_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.template_folder = folder_path
            self.save_setting(folder_path)
            self.load_template_checkboxes()

    def save_setting(self, path):
        setting_file = 'setting.txt'
        try:
            with open(setting_file, 'w') as file:
                file.write(path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save settings: {str(e)}")

    def load_setting(self):
        setting_file = 'setting.txt'
        if not os.path.exists(setting_file):
            with open(setting_file, 'w') as file:
                file.write("")
            messagebox.showinfo("Info", "Setting file not found, created a new one.")
        else:
            with open(setting_file, 'r') as file:
                self.template_folder = file.read().strip()
            self.load_template_checkboxes()

    def load_template_checkboxes(self):
        for checkbox in self.checkboxes:
            checkbox.destroy()
        self.checkboxes = []

        if self.template_folder:
            for i, filename in enumerate(os.listdir(self.template_folder)):
                if filename.endswith('.docx'):
                    var = tk.IntVar()
                    checkbox = tk.Checkbutton(self.checkboxes_frame, text=filename, variable=var, command=lambda i=i, var=var: self.on_checkbox_toggle(i, filename, var))
                    checkbox.var = var
                    checkbox.pack(anchor='w')
                    self.checkboxes.append(checkbox)

    def on_checkbox_toggle(self, index, name, var):
        print(f"Checkbox #{index + 1} - {name} - {'Selected' if var.get() else 'Deselected'}")

    def add_row(self):
        self.table.insert('', 'end', values=('', ''))

    def remove_row(self):
        selected_item = self.table.selection()
        if selected_item:
            self.table.delete(selected_item)

    def save_table(self):
        if not self.file_name_label.cget("text"):
            messagebox.showwarning("Warning", "No file selected to save the table.")
            return

        table_data = []
        for row_id in self.table.get_children():
            row = self.table.item(row_id)['values']
            table_data.append(row)

        file_name = os.path.splitext(self.file_name_label.cget("text"))[0] + '.txt'
        try:
            with open(file_name, 'w', encoding='utf-8') as file:
                for row in table_data:
                    file.write('\t'.join(map(str, row)) + '\n')
            messagebox.showinfo("Info", f"Table saved as {file_name}")
        except IOError as e:
            messagebox.showerror("Error", f"Failed to save table: {str(e)}")
        except Exception as e:
            messagebox.showerror("Error", f"Unexpected error: {str(e)}")

    def checkall(self):
        current_state = self.check_all_var.get()
        new_state = not current_state
        self.check_all_var.set(new_state)

        for checkbox in self.checkboxes:
            checkbox.var.set(new_state)
            self.on_checkbox_toggle(self.checkboxes.index(checkbox), checkbox.cget('text'), checkbox.var)

    def edit_process(self):
        backend = Backend(self.file_path, "")
        backend.run()
        i=0
        while i <= len(backend.row_data) - 1 :
            self.table.insert('', 'end', values=(backend.headers[i], backend.row_data[i]))
            i=i+1

    def start_process(self):
        self.end_data = []
        if not self.file_name_label.cget("text"):
            messagebox.showwarning("Warning", "No file selected to save the table.")
            return

        table_data = []
        for row_id in self.table.get_children():
            row = self.table.item(row_id)['values']
            table_data.append(row)

        file_name = os.path.splitext(self.file_name_label.cget("text"))[0] + '.txt'
        with open(file_name, 'w', encoding='utf-8') as file:
            for row in table_data:
                self.end_data.append(row)

        checkbox_info = {}
        true_checkboxes = {}

        for i, checkbox in enumerate(self.checkboxes):
            name = checkbox.cget('text')
            state = bool(checkbox.var.get())
            checkbox_info[f'checkbox#{i + 1}'] = {'name': name, 'state': state}
            if state:
                true_checkboxes[f'checkbox#{i + 1}'] = {'name': name, 'state': state}

        print(true_checkboxes)
        print(self.end_data)

        # Чтение пути к шаблонам из setting.txt с учетом возможных проблем с кодировкой
        try:
            with open('setting.txt', 'r', encoding='cp1251') as file:
                template_path = file.readline().strip()
        except UnicodeDecodeError:
            messagebox.showerror("Error", "Failed to read 'setting.txt'. Please check the file encoding.")
            return
        except FileNotFoundError:
            messagebox.showerror("Error", "'setting.txt' not found.")
            return

        # Создание папки, если её нет
        dest_folder = self.file_name_label.cget("text")
        if not os.path.exists(dest_folder):
            os.makedirs(dest_folder)

        # Копирование файлов и замена слов
        for checkbox in true_checkboxes.values():
            source_file = os.path.join(template_path, checkbox['name'])
            dest_file = os.path.join(dest_folder, checkbox['name'])
            shutil.copy2(source_file, dest_file)

            # Открытие файла и замена слов
            self.replace_words_in_file(dest_file, self.end_data)

    def replace_words_in_file(self, file_path, replacements):
        try:
            doc = Document(file_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open {file_path}: {e}")
            return

        for paragraph in doc.paragraphs:
            for key, value in replacements:
                paragraph.text = paragraph.text.replace(f"({key})", str(value))  # Преобразуем второй аргумент в строку

        try:
            doc.save(file_path)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save {file_path}: {e}")


class Backend:
    def __init__(self, input_path, output_path):
        self.input_path = input_path
        self.output_path = output_path
        self.data = {}
        self.morph = pymorphy2.MorphAnalyzer()  # Initialize pymorphy2.MorphAnalyzer

    def extract_data(self, text):
        data = {}

        # Номер документа
        match_numb = re.search(r'No(\d+) ФГБУ', text)
        if match_numb:
            data["Numb"] = match_numb.group(1).strip()

        # Extract rank, part, and name
        match_rank_part = re.search(r'Рядовой в/ч (\d+) (.*?)\n', text)
        if match_rank_part:
            data["Rank"] = "Рядовой"
            data["Part"] = match_rank_part.group(1).strip()
            data["Fio"] = match_rank_part.group(2).strip()

        # Extract date of birth
        match_dob = re.search(r'Дата рождения (\d{2}\.\d{2}\.\d{4}) г.р.', text)
        if match_dob:
            data["DOB"] = match_dob.group(1).strip()

        # Extract admission and discharge dates
        match_dates = re.search(r'Находился на лечении с (\d{2}\.\d{2}\.\d{4}) г. по (\d{2}\.\d{2}\.\d{4}) г.', text)
        if match_dates:
            data["Admission"] = match_dates.group(1).strip()
            data["Discharge"] = match_dates.group(2).strip()

        # Extract diagnosis
        match_diagnosis = re.search(r'Основной диагноз: (.*?)\n', text)
        if match_diagnosis:
            data["Diagnosis"] = match_diagnosis.group(1).strip()

        # Extract complaints
        match_complaints = re.search(r'Жалобы: (.*?)\n', text)
        if match_complaints:
            data["Complaints"] = match_complaints.group(1).strip()

        # Extract anamnesis
        match_anamnesis = re.search(r'Анамнез: (.*?)\n', text)
        if match_anamnesis:
            data["Anamnesis"] = match_anamnesis.group(1).strip()

        # Extract objective status
        match_objective_status = re.search(r'Объективный статус: (.*?)\n', text)
        if match_objective_status:
            data["ObjectiveStatus"] = match_objective_status.group(1).strip()

        # Extract cardiovascular system status
        match_cvs = re.search(r'Сердечно-сосудистая система: (.*?)\n', text)
        if match_cvs:
            data["CardioVascularSystem"] = match_cvs.group(1).strip()

        # Extract respiratory system status
        match_respiratory_system = re.search(r'Дыхательная система: (.*?)\n', text)
        if match_respiratory_system:
            data["RespiratorySystem"] = match_respiratory_system.group(1).strip()

        # Extract digestive system status
        match_digestive_system = re.search(r'Система органов пищеварения: (.*?)\n', text)
        if match_digestive_system:
            data["DigestiveSystem"] = match_digestive_system.group(1).strip()

        # Extract urinary system status
        match_urinary_system = re.search(r'Мочеполовая система: (.*?)\n', text)
        if match_urinary_system:
            data["Urinary"] = match_urinary_system.group(1).strip()

        # Extract laboratory data
        match_laboratory_data = re.search(r'Данные лабораторных методов исследования: (.*?)\n', text)
        if match_laboratory_data:
            data["LaboratoryData"] = match_laboratory_data.group(1).strip()

        # Extract therapy data
        match_therapy = re.search(r'Получал терапию: (.*?)\n', text)
        if match_therapy:
            data["Therapy"] = match_therapy.group(1).strip()

        # Extract instrumental methods data (tools)
        match_tools = re.search(r'Данные инструментальных методов исследования: (.*?)\n', text)
        if match_tools:
            data["Tools"] = match_tools.group(1).strip()

        # Extract diagnosismakeup (genitive case diagnosis)
        match_diagnosismakeup = re.search(r'Основной диагноз: (.*?)\n', text)
        if match_diagnosismakeup:
            diagnosismakeup = match_diagnosismakeup.group(1).strip()

            # Split diagnosismakeup into words if needed
            words = diagnosismakeup.split()

            # Convert each word to genitive case
            genitive_words = []
            for word in words:
                parsed_word = self.morph.parse(word)[0]  # get the first parsing result
                genitive_form = parsed_word.inflect({'gent'}).word
                genitive_words.append(genitive_form)

            # Join the words back into a string
            data["Diagnosismakeup"] = ' '.join(genitive_words)

        return data

    def process_docx(self):
        # Load the DOCX file
        doc = Document(self.input_path)
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Extract data from the text
        self.data = self.extract_data(full_text)

    def save_to_xlsx(self):
        # Write extracted data to XLSX file
        wb = Workbook()
        ws = wb.active

        # Define headers based on your description
        self.headers = ["Numb", "Fio", "Part", "Rank", "Agedate", "Receiveddate", "Pity", "Information", "Condition",
                   "Vessels", "Breath", "Digestion", "Urinary", "Laboratory", "Tools", "Therapy",
                   "Diagnosismakeup", "Diagnosis"]

        # Write headers to the first row
        ws.append(self.headers)

        # Prepare data row
        self.row_data = [
            self.data.get("Numb", "-"),  # Numb
            self.data.get("Fio", ""),  # Fio
            self.data.get("Part", ""),  # Part
            self.data.get("Rank", ""),  # Rank
            self.data.get("DOB", ""),  # Agedate
            self.data.get("Admission", ""),  # Receiveddate
            self.data.get("Complaints", ""),  # Pity
            self.data.get("Anamnesis", ""),  # Information
            self.data.get("ObjectiveStatus", ""),  # Condition
            self.data.get("CardioVascularSystem", ""),  # Vessels
            self.data.get("RespiratorySystem", ""),  # Breath
            self.data.get("DigestiveSystem", ""),  # Digestion
            self.data.get("Urinary", ""),  # Urinary
            self.data.get("LaboratoryData", ""),  # Laboratory
            self.data.get("Tools", ""),  # Tools
            self.data.get("Therapy", ""),  # Therapy
            self.data.get("Diagnosismakeup", ""),  # Diagnosismakeup
            self.data.get("Diagnosis", "")  # Diagnosis
        ]

        # Append the data row to the worksheet
        ws.append(self.row_data)



    def run(self):
        self.process_docx()
        self.save_to_xlsx()


if __name__ == "__main__":
    root = tk.Tk()
    app = GUITestApp(root)
    app.load_setting()
    root.mainloop()
